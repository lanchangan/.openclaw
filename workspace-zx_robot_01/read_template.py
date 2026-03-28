#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""读取Word模板，分析其排版结构"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

template_path = r'C:\Users\LWS\.openclaw\workspace-zx_robot_01\糟雪_教师岗.docx'
doc = Document(template_path)

print("=== Word文档结构分析 ===\n")
print(f"总段落数: {len(doc.paragraphs)}\n")

for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip():
        style = paragraph.style
        alignment = paragraph.alignment
        indent = paragraph.paragraph_format.left_indent
        space_before = paragraph.paragraph_format.space_before
        space_after = paragraph.paragraph_format.space_after
        
        print(f"段落 {i+1}:")
        print(f"  文本: {repr(paragraph.text)}")
        print(f"  样式: {style.name}")
        print(f"  对齐: {alignment}")
        print(f"  左缩进: {indent}")
        print(f"  段前间距: {space_before}")
        print(f"  段后间距: {space_after}")
        print()

print("\n=== 表格数量 ===")
print(f"表格数: {len(doc.tables)}")
