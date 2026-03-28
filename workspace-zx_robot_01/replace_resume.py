#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于现有Word模板替换简历内容
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 简历内容
resume_content = {
    'name': '糟雪',
    'birth': '1998年5月2日',
    'job_intent': '小学语文教师',
    'phone': '',
    'email': '',
    'education': [
        {'time': '2020.09—2023.07', 'school': '西北大学', 'major': '中国古代文学专业', 'degree': '硕士研究生'},
        {'time': '2016.09—2020.06', 'school': '西北师范大学', 'major': '汉语言文学专业', 'degree': '本科'}
    ],
    'work': [
        {
            'time': '2023.09—至今',
            'company': '西北大学附属小学',
            'position': '语文教师兼班主任',
            'points': [
                '全面负责小学语文教学与班主任工作，所带班级成绩优异、稳居年级第一，连续获得学校教学质量奖，班级多次获评校级先进班集体，本人荣获校级优秀教师称号。',
                '扎实开展班级管理与家校共育工作，注重学生习惯养成与品格培养，班风学风优良，获得学校、家长高度认可。',
                '作为学校语文教学核心骨干，深度参与校级教研、课程研发与课题研究，在校长名师工作室中承担重要课题任务。',
                '擅长AI赋能语文教学，在智慧课堂与数字化教学中发挥引领作用，独立开发《守株待兔》《海底世界》《夏日绝句》等优质AI赋能语文教学课例。',
                '课程论文每次均荣获校级一等奖第一名，AI赋能教学成果在校内及碑林区内示范推广，教学创新能力突出。'
            ]
        }
    ],
    'internship': [
        {'time': '2023.05—2023.06', 'school': '西安建筑科技大学附属小学', 'position': '语文教师实习'},
        {'time': '2019.09—2019.11', 'school': '临夏回民中学', 'position': '语文教师实习'}
    ],
    'skills': [
        '教师资格：高级中学语文教师资格证',
        '语言水平：普通话二级甲等',
        '教学能力：精通小学语文全学段教学、班级管理、AI智慧教学、课例研发与课程论文撰写',
        '教研能力：具备区级教学成果推广经验、名师工作室课题研究经验'
    ],
    'self_evaluation': '本人本硕均为中文专业，文学功底扎实，专业素养深厚。拥有丰富的小学语文教学与班主任工作经验，教学成绩优异，班级管理成果显著。擅长AI赋能语文教学，教研能力突出，课例与论文多次获得校级最高奖项，教学成果在区域内示范推广。热爱教育事业，责任心强，亲和力佳，具备优秀的教学、教研与班级管理能力，能够快速胜任小学语文教师岗位。'
}

def create_new_resume(template_path, output_path):
    """基于模板创建新简历"""
    doc = Document(template_path)
    
    # 清除原有内容，保留样式
    for paragraph in doc.paragraphs:
        paragraph.clear()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    
    # 标题 - 姓名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(resume_content['name'])
    run.bold = True
    run.font.size = Pt(22)
    
    # 基本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    info_text = f"出生年月：{resume_content['birth']}  求职意向：{resume_content['job_intent']}"
    if resume_content['phone']:
        info_text += f"\n联系电话：{resume_content['phone']}"
    if resume_content['email']:
        info_text += f"  电子邮箱：{resume_content['email']}"
    run = p.add_run(info_text)
    run.font.size = Pt(11)
    
    # 教育背景标题
    p = doc.add_paragraph()
    run = p.add_run('教育背景')
    run.bold = True
    run.font.size = Pt(13)
    p.space_before = Pt(6)
    p.space_after = Pt(3)
    
    # 教育背景内容
    for edu in resume_content['education']:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        text = f"{edu['time']}  {edu['school']}  {edu['major']}  {edu['degree']}"
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    
    # 工作经历标题
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(3)
    run = p.add_run('工作经历')
    run.bold = True
    run.font.size = Pt(13)
    
    # 工作经历内容
    for work in resume_content['work']:
        p = doc.add_paragraph()
        p.space_after = Pt(3)
        text = f"{work['time']}  {work['company']}  {work['position']}"
        run = p.add_run(text)
        run.bold = False
        run.font.size = Pt(10.5)
        
        for point in work['points']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.space_after = Pt(2)
            run = p.add_run(f"• {point}")
            run.font.size = Pt(10.5)
    
    # 实习经历标题
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(3)
    run = p.add_run('实习经历')
    run.bold = True
    run.font.size = Pt(13)
    
    # 实习经历内容
    for intern in resume_content['internship']:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        text = f"{intern['time']}  {intern['school']}  {intern['position']}"
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    
    # 专业技能与证书标题
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(3)
    run = p.add_run('专业技能与证书')
    run.bold = True
    run.font.size = Pt(13)
    
    # 专业技能内容
    for skill in resume_content['skills']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.space_after = Pt(2)
        run = p.add_run(f"• {skill}")
        run.font.size = Pt(10.5)
    
    # 自我评价标题
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(3)
    run = p.add_run('自我评价')
    run.bold = True
    run.font.size = Pt(13)
    
    # 自我评价内容
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(resume_content['self_evaluation'])
    run.font.size = Pt(10.5)
    
    # 保存文档
    doc.save(output_path)
    print(f"简历已生成：{output_path}")

if __name__ == '__main__':
    template_path = r'C:\Users\LWS\.openclaw\workspace-zx_robot_01\糟雪_教师岗.docx'
    output_path = r'C:\Users\LWS\.openclaw\workspace-zx_robot_01\糟雪_小学语文教师简历.docx'
    create_new_resume(template_path, output_path)
