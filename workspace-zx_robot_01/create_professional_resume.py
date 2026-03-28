#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建专业排版的小学语文教师简历
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)

# 标题：姓名
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('糟  雪')
run.bold = True
run.font.size = Pt(26)
p.space_after = Pt(6)

# 基本信息
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(6)
run = p.add_run('出生年月：1998年5月  求职意向：小学语文教师')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(12)
run = p.add_run('联系电话：__________  电子邮箱：____________________')
run.font.size = Pt(11)

# 添加分隔线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 60)
run.font.size = Pt(2)
p.space_after = Pt(6)

# 教育背景
p = doc.add_paragraph()
p.space_before = Pt(3)
p.space_after = Pt(6)
run = p.add_run('教育背景')
run.bold = True
run.font.size = Pt(14)
run.underline = True

# 教育背景内容
p = doc.add_paragraph()
p.space_after = Pt(2)
run = p.add_run('■ 2020.09 — 2023.07  西北大学  中国古代文学专业  硕士研究生')
run.font.size = Pt(10.5)
run.bold = False

p = doc.add_paragraph()
p.space_after = Pt(6)
run = p.add_run('■ 2016.09 — 2020.06  西北师范大学  汉语言文学专业  本科')
run.font.size = Pt(10.5)
run.bold = False

# 工作经历
p = doc.add_paragraph()
p.space_before = Pt(6)
p.space_after = Pt(6)
run = p.add_run('工作经历')
run.bold = True
run.font.size = Pt(14)
run.underline = True

p = doc.add_paragraph()
p.space_after = Pt(3)
run = p.add_run('■ 2023.09 — 至今  西北大学附属小学  语文教师兼班主任')
run.font.size = Pt(10.5)
run.bold = True

# 工作内容要点
work_points = [
    '全面负责小学语文教学与班主任工作，所带班级成绩优异、稳居年级第一，连续获得学校教学质量奖，班级多次获评校级先进班集体，本人荣获校级优秀教师称号。',
    '扎实开展班级管理与家校共育工作，注重学生习惯养成与品格培养，班风学风优良，获得学校、家长高度认可。',
    '作为学校语文教学核心骨干，深度参与校级教研、课程研发与课题研究，在校长名师工作室中承担重要课题任务。',
    '擅长AI赋能语文教学，在智慧课堂与数字化教学中发挥引领作用，独立开发《守株待兔》《海底世界》《夏日绝句》等优质AI赋能语文教学课例。',
    '课程论文每次均荣获校级一等奖第一名，AI赋能教学成果在校内及碑林区内示范推广，教学创新能力突出。'
]

for point in work_points:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.space_after = Pt(2)
    run = p.add_run('•  ' + point)
    run.font.size = Pt(10.5)

# 实习经历
p = doc.add_paragraph()
p.space_before = Pt(6)
p.space_after = Pt(6)
run = p.add_run('实习经历')
run.bold = True
run.font.size = Pt(14)
run.underline = True

p = doc.add_paragraph()
p.space_after = Pt(2)
run = p.add_run('■ 2023.05 — 2023.06  西安建筑科技大学附属小学  语文教师实习')
run.font.size = Pt(10.5)

p = doc.add_paragraph()
p.space_after = Pt(6)
run = p.add_run('■ 2019.09 — 2019.11  临夏回民中学  语文教师实习')
run.font.size = Pt(10.5)

# 专业技能与证书
p = doc.add_paragraph()
p.space_before = Pt(6)
p.space_after = Pt(6)
run = p.add_run('专业技能与证书')
run.bold = True
run.font.size = Pt(14)
run.underline = True

skills = [
    '▫️ 教师资格：高级中学语文教师资格证',
    '▫️ 语言水平：普通话二级甲等',
    '▫️ 教学能力：精通小学语文全学段教学、班级管理、AI智慧教学、课例研发与课程论文撰写',
    '▫️ 教研能力：具备区级教学成果推广经验、名师工作室课题研究经验'
]

for skill in skills:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.space_after = Pt(3)
    run = p.add_run(skill)
    run.font.size = Pt(10.5)

# 自我评价
p = doc.add_paragraph()
p.space_before = Pt(6)
p.space_after = Pt(6)
run = p.add_run('自我评价')
run.bold = True
run.font.size = Pt(14)
run.underline = True

p = doc.add_paragraph()
p.space_after = Pt(6)
run = p.add_run('本人本硕均为中文专业，文学功底扎实，专业素养深厚。拥有丰富的小学语文教学与班主任工作经验，教学成绩优异，班级管理成果显著。擅长AI赋能语文教学，教研能力突出，课例与论文多次获得校级最高奖项，教学成果在区域内示范推广。热爱教育事业，责任心强，亲和力佳，具备优秀的教学、教研与班级管理能力，能够快速胜任小学语文教师岗位。')
run.font.size = Pt(10.5)

# 保存文档
output_path = r'C:\Users\LWS\.openclaw\workspace-zx_robot_01\糟雪_小学语文教师简历_专业排版.docx'
doc.save(output_path)
print(f"专业排版简历已生成: {output_path}")
