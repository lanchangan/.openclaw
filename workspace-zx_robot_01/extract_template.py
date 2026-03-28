#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""提取Word模板的完整内容和样式"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

template_path = r'C:\Users\LWS\.openclaw\workspace-zx_robot_01\糟雪_教师岗.docx'
doc = Document(template_path)

print("=== 完整文档内容 ===\n")

all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text)
        print(f"{para.text}")

print(f"\n总共 {len(all_text)} 个非空段落")
print(f"\n文档有 {len(doc.tables)} 个表格")

# 打印样式信息
print("\n=== 样式列表 ===")
for style in doc.styles:
    if style.type == 1:  # Paragraph styles
        if style.font.size:
            print(f"{style.name}: 字号 {style.font.size}")
