#!/usr/bin/env python3
"""
生成Word文档 - Bilibili视频逐字稿
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold

def create_word_document():
    # 读取逐字稿内容
    transcript_path = 'bili_temp/BV1u61mBhEj7/BV1u61mBhEj7_transcript.txt'
    
    with open(transcript_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('', 0)
    title_run = title.add_run('四年级下册《看看我们的地球》导读课')
    set_chinese_font(title_run, 'Microsoft YaHei', 18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('讲师：王秀萍')
    set_chinese_font(subtitle_run, 'Microsoft YaHei', 12, bold=False)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分隔线
    doc.add_paragraph()
    
    # 添加逐字稿标题
    section_title = doc.add_paragraph()
    section_run = section_title.add_run('【课堂逐字稿】')
    set_chinese_font(section_run, 'Microsoft YaHei', 14, bold=True)
    
    doc.add_paragraph()
    
    # 解析并添加逐字稿内容
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 跳过标题行
        if line.startswith('视频标题:') or line.startswith('讲师:') or line.startswith('BV号:'):
            continue
        
        # 跳过分隔线和标题
        if line.startswith('=') or line == '【课堂逐字稿】':
            continue
        
        # 处理时间戳和文本
        if line.startswith('[') and ']' in line:
            # 提取时间戳和内容
            match = re.match(r'\[(\d{2}:\d{2})\]\s*(.+)', line)
            if match:
                timestamp = match.group(1)
                text = match.group(2)
                
                # 创建段落
                p = doc.add_paragraph()
                
                # 添加时间戳（灰色）
                time_run = p.add_run(f'[{timestamp}] ')
                set_chinese_font(time_run, 'SimSun', 10)
                time_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # 添加内容
                text_run = p.add_run(text)
                set_chinese_font(text_run, 'SimSun', 11)
    
    # 添加结尾
    doc.add_paragraph()
    end_para = doc.add_paragraph()
    end_run = end_para.add_run('—— 课程结束 ——')
    set_chinese_font(end_run, 'Microsoft YaHei', 12)
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    output_path = '四年级下册《看看我们的地球》导读课-逐字稿.docx'
    doc.save(output_path)
    print(f'✅ Word文档已生成: {output_path}')

if __name__ == '__main__':
    create_word_document()
