from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 添加标题
title = doc.add_heading('四年级下册《看看我们的地球》导读课', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加副标题
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('王秀萍 老师')
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# 读取字幕内容
with open('bili_temp/BV1u61mBhEj7/BV1u61mBhEj7_chunk_0.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# 添加正文内容
doc.add_heading('课堂逐字稿', 1)

# 将内容分段处理
paragraphs = content.split('\n\n')

for para_text in paragraphs:
    para_text = para_text.strip()
    if para_text:
        # 检查是否是时间戳行
        if para_text.startswith('[') and ']' in para_text:
            # 提取时间戳和文本
            time_end = para_text.find(']')
            timestamp = para_text[:time_end+1]
            text_content = para_text[time_end+1:].strip()
            
            p = doc.add_paragraph()
            time_run = p.add_run(timestamp + ' ')
            time_run.font.size = Pt(10)
            time_run.font.color.rgb = RGBColor(150, 150, 150)
            p.add_run(text_content)
        else:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.line_spacing = 1.15

# 添加结尾
doc.add_paragraph()
doc.add_paragraph('—— 课程结束 ——').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_path = '四年级下册《看看我们的地球》导读课-逐字稿.docx'
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
