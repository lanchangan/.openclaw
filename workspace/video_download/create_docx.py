# -*- coding: utf-8 -*-
from docx import Documentnfrom docx.shared import Pt, Inches
# 读取 transcriptnwith open('四年级下册《看看我们的地球》导读课-王秀萍.txt', 'r', encoding='utf-8') as f:n    content = f.read()nn# 创建 Word docdoc = Document()n
# 添加标题ntitle = doc.add_heading('四年级下册《看看我们的地球》导读课-王秀萍', 0)ntitle.alignment = WD_ALIGN_CENTERn
# 添加正文paragraphs = content.split('\n')nfor line in paragraphs:n    if line.strip():  # 跳过空行n        p = doc.add_paragraph(line)
        p.style = 'List Bullet'
# 保存doc.save('四年级下册《看看我们的地球》导读课-王秀萍-逐字稿.docx')
print('Word文档已保存完毕！')